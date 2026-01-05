"""
NORMAEX 2.0 - Document Executor
Sistema de execução de ações de formatação baseado em planos da IA
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any
import os


class DocumentExecutor:
    """
    Executor de ações de formatação no documento
    """

    def __init__(self, doc_path: str):
        """
        Inicializa o executor com um documento

        Args:
            doc_path: Caminho para o documento DOCX
        """
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"Documento não encontrado: {doc_path}")

        self.doc = Document(doc_path)
        self.doc_path = doc_path
        self.actions_log = []

    def execute_action_plan(self, action_plan: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Executa todas as ações do plano sequencialmente

        Args:
            action_plan: Lista de ações a executar

        Returns:
            list: Resultados de cada ação
        """
        results = []

        # Ordenar por prioridade
        sorted_plan = sorted(action_plan, key=lambda x: x.get("priority", 999))

        for idx, action in enumerate(sorted_plan):
            try:
                result = self.execute_single_action(action)
                results.append({
                    "action_index": idx,
                    "action": action,
                    "status": "success",
                    "result": result,
                    "message": f"✓ {action.get('description', 'Ação executada')}"
                })

                # Adicionar ao log
                self.actions_log.append({
                    "action": action["action"],
                    "target": action.get("target"),
                    "status": "success",
                    "result": result
                })

            except Exception as e:
                results.append({
                    "action_index": idx,
                    "action": action,
                    "status": "error",
                    "error": str(e),
                    "message": f"✗ Erro: {str(e)}"
                })

                # Adicionar ao log
                self.actions_log.append({
                    "action": action["action"],
                    "target": action.get("target"),
                    "status": "error",
                    "error": str(e)
                })

        return results

    def execute_single_action(self, action: Dict[str, Any]) -> str:
        """
        Executa uma ação individual

        Args:
            action: Dicionário com tipo de ação e parâmetros

        Returns:
            str: Mensagem de resultado
        """
        action_type = action["action"]

        if action_type == "fix_margin":
            return self.fix_margin(action["target"], action["params"])

        elif action_type == "fix_font":
            return self.fix_font(action["target"], action["params"])

        elif action_type == "fix_spacing":
            return self.fix_spacing(action["target"], action["params"])

        elif action_type == "fix_alignment":
            return self.fix_alignment(action["target"], action["params"])

        elif action_type == "fix_indent":
            return self.fix_indent(action["target"], action["params"])

        else:
            raise ValueError(f"Ação desconhecida: {action_type}")

    def fix_margin(self, target: str, params: Dict[str, float]) -> str:
        """
        Corrige margens de uma seção

        Args:
            target: Identificador da seção (ex: "section_0")
            params: Margens em cm (top, bottom, left, right)

        Returns:
            str: Mensagem de resultado
        """
        if target.startswith("section_"):
            section_idx = int(target.split("_")[1])

            if section_idx >= len(self.doc.sections):
                raise ValueError(f"Seção {section_idx} não existe no documento")

            section = self.doc.sections[section_idx]

            # Aplicar margens
            if "top" in params:
                section.top_margin = Cm(params["top"])
            if "bottom" in params:
                section.bottom_margin = Cm(params["bottom"])
            if "left" in params:
                section.left_margin = Cm(params["left"])
            if "right" in params:
                section.right_margin = Cm(params["right"])

            return f"Margens ajustadas na seção {section_idx}: {params}"

        elif target == "all_sections":
            count = 0
            for section in self.doc.sections:
                if "top" in params:
                    section.top_margin = Cm(params["top"])
                if "bottom" in params:
                    section.bottom_margin = Cm(params["bottom"])
                if "left" in params:
                    section.left_margin = Cm(params["left"])
                if "right" in params:
                    section.right_margin = Cm(params["right"])
                count += 1

            return f"Margens ajustadas em {count} seções: {params}"

        else:
            raise ValueError(f"Target inválido para fix_margin: {target}")

    def fix_font(self, target: str, params: Dict[str, Any]) -> str:
        """
        Corrige fonte dos parágrafos

        Args:
            target: Identificador do alvo (ex: "all_body", "paragraph_5")
            params: Propriedades de fonte (name, size)

        Returns:
            str: Mensagem de resultado
        """
        font_name = params.get("name", "Arial")
        font_size = params.get("size", 12)

        if target == "all_body":
            count = 0
            for para in self.doc.paragraphs:
                # Pular títulos (headings)
                if not para.style.name.lower().startswith('heading') and \
                   not para.style.name.lower().startswith('título'):
                    for run in para.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                    count += 1

            return f"Fonte ajustada para {font_name} {font_size}pt em {count} parágrafos"

        elif target.startswith("paragraph_"):
            para_idx = int(target.split("_")[1])

            if para_idx >= len(self.doc.paragraphs):
                raise ValueError(f"Parágrafo {para_idx} não existe no documento")

            para = self.doc.paragraphs[para_idx]
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)

            return f"Fonte ajustada no parágrafo {para_idx}: {font_name} {font_size}pt"

        elif target == "all":
            count = 0
            for para in self.doc.paragraphs:
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                count += 1

            return f"Fonte ajustada em {count} parágrafos (incluindo títulos)"

        else:
            raise ValueError(f"Target inválido para fix_font: {target}")

    def fix_spacing(self, target: str, params: Dict[str, float]) -> str:
        """
        Corrige espaçamento entre linhas

        Args:
            target: Identificador do alvo
            params: Espaçamento (line_spacing)

        Returns:
            str: Mensagem de resultado
        """
        line_spacing = params.get("line_spacing", 1.5)

        if target == "all_body":
            count = 0
            for para in self.doc.paragraphs:
                # Pular títulos
                if not para.style.name.lower().startswith('heading') and \
                   not para.style.name.lower().startswith('título'):
                    para.paragraph_format.line_spacing = line_spacing
                    count += 1

            return f"Espaçamento ajustado para {line_spacing} em {count} parágrafos"

        elif target.startswith("paragraph_"):
            para_idx = int(target.split("_")[1])

            if para_idx >= len(self.doc.paragraphs):
                raise ValueError(f"Parágrafo {para_idx} não existe no documento")

            para = self.doc.paragraphs[para_idx]
            para.paragraph_format.line_spacing = line_spacing

            return f"Espaçamento ajustado no parágrafo {para_idx}: {line_spacing}"

        elif target == "all":
            count = 0
            for para in self.doc.paragraphs:
                para.paragraph_format.line_spacing = line_spacing
                count += 1

            return f"Espaçamento ajustado em {count} parágrafos (incluindo títulos)"

        else:
            raise ValueError(f"Target inválido para fix_spacing: {target}")

    def fix_alignment(self, target: str, params: Dict[str, str]) -> str:
        """
        Corrige alinhamento dos parágrafos

        Args:
            target: Identificador do alvo
            params: Alinhamento (alignment: "justify", "left", "center", "right")

        Returns:
            str: Mensagem de resultado
        """
        alignment_map = {
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }

        alignment_str = params.get("alignment", "justify")
        alignment = alignment_map.get(alignment_str, WD_ALIGN_PARAGRAPH.JUSTIFY)

        if target == "all_body":
            count = 0
            for para in self.doc.paragraphs:
                # Pular títulos
                if not para.style.name.lower().startswith('heading') and \
                   not para.style.name.lower().startswith('título'):
                    para.alignment = alignment
                    count += 1

            return f"Alinhamento ajustado para '{alignment_str}' em {count} parágrafos"

        elif target.startswith("paragraph_"):
            para_idx = int(target.split("_")[1])

            if para_idx >= len(self.doc.paragraphs):
                raise ValueError(f"Parágrafo {para_idx} não existe no documento")

            para = self.doc.paragraphs[para_idx]
            para.alignment = alignment

            return f"Alinhamento ajustado no parágrafo {para_idx}: {alignment_str}"

        elif target == "all":
            count = 0
            for para in self.doc.paragraphs:
                para.alignment = alignment
                count += 1

            return f"Alinhamento ajustado em {count} parágrafos (incluindo títulos)"

        else:
            raise ValueError(f"Target inválido para fix_alignment: {target}")

    def fix_indent(self, target: str, params: Dict[str, float]) -> str:
        """
        Corrige recuo da primeira linha

        Args:
            target: Identificador do alvo
            params: Recuo (first_line em cm)

        Returns:
            str: Mensagem de resultado
        """
        first_line_indent = params.get("first_line", 1.25)

        if target == "all_body":
            count = 0
            for para in self.doc.paragraphs:
                # Pular títulos e parágrafos vazios
                if not para.style.name.lower().startswith('heading') and \
                   not para.style.name.lower().startswith('título') and \
                   para.text.strip():
                    para.paragraph_format.first_line_indent = Cm(first_line_indent)
                    count += 1

            return f"Recuo ajustado para {first_line_indent}cm em {count} parágrafos"

        elif target.startswith("paragraph_"):
            para_idx = int(target.split("_")[1])

            if para_idx >= len(self.doc.paragraphs):
                raise ValueError(f"Parágrafo {para_idx} não existe no documento")

            para = self.doc.paragraphs[para_idx]
            para.paragraph_format.first_line_indent = Cm(first_line_indent)

            return f"Recuo ajustado no parágrafo {para_idx}: {first_line_indent}cm"

        elif target == "all":
            count = 0
            for para in self.doc.paragraphs:
                if para.text.strip():  # Apenas parágrafos não vazios
                    para.paragraph_format.first_line_indent = Cm(first_line_indent)
                    count += 1

            return f"Recuo ajustado em {count} parágrafos (incluindo títulos)"

        else:
            raise ValueError(f"Target inválido para fix_indent: {target}")

    def save(self, output_path: str) -> str:
        """
        Salva documento modificado

        Args:
            output_path: Caminho de saída

        Returns:
            str: Mensagem de confirmação
        """
        try:
            self.doc.save(output_path)
            return f"Documento salvo com sucesso: {output_path}"
        except Exception as e:
            raise Exception(f"Erro ao salvar documento: {str(e)}")

    def get_actions_log(self) -> List[Dict[str, Any]]:
        """
        Retorna log de todas as ações executadas

        Returns:
            list: Log de ações
        """
        return self.actions_log

    def get_summary(self) -> Dict[str, Any]:
        """
        Retorna resumo da execução

        Returns:
            dict: Resumo com estatísticas
        """
        total_actions = len(self.actions_log)
        successful_actions = len([a for a in self.actions_log if a["status"] == "success"])
        failed_actions = len([a for a in self.actions_log if a["status"] == "error"])

        return {
            "total_actions": total_actions,
            "successful_actions": successful_actions,
            "failed_actions": failed_actions,
            "success_rate": (successful_actions / total_actions * 100) if total_actions > 0 else 0,
            "actions_log": self.actions_log
        }
