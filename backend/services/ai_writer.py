"""
Normaex 2.0 - FASE 5: Escrita Inteligente
Escritor de IA que gera texto JÁ formatado estruturalmente
"""

import os
import json
import google.generativeai as genai
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.executor import DocumentExecutor


# Configurar API do Gemini (opcional no import)
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)


def write_with_structure(
    document_context: str,
    instruction: str,
    section_type: str,
    document_structure: Dict[str, Any]
) -> Dict[str, Any]:
    """
    IA escreve texto já formatado estruturalmente

    Args:
        document_context: Texto do documento atual
        instruction: Instrução do usuário sobre o que escrever
        section_type: Tipo da seção (introducao, metodologia, etc.)
        document_structure: Estrutura completa do documento

    Returns:
        dict: {
            "content": "Texto gerado",
            "structure": {
                "type": "body",
                "formatting": {...},
                "action_plan": [...]
            }
        }
    """
    # Verificar se API key está configurada
    if not GOOGLE_API_KEY:
        raise ValueError("GOOGLE_API_KEY não configurada. Configure a variável de ambiente.")

    model = genai.GenerativeModel("gemini-2.5-flash")

    # Preparar resumo da estrutura (evitar JSON muito grande)
    structure_summary = {
        "total_paragraphs": len(document_structure.get("paragraphs", [])),
        "total_sections": len(document_structure.get("sections", [])),
        "hierarchy": document_structure.get("hierarchy", []),
        "last_paragraph_index": len(document_structure.get("paragraphs", [])) - 1
    }

    prompt = f"""
Você é um escritor acadêmico especialista em normas ABNT.

CONTEXTO DO DOCUMENTO:
```
{document_context[-2000:]}  # Últimos 2000 caracteres para contexto
```

ESTRUTURA DO DOCUMENTO:
```json
{json.dumps(structure_summary, indent=2, ensure_ascii=False)}
```

INSTRUÇÃO DO USUÁRIO:
{instruction}

TIPO DE SEÇÃO: {section_type}

TAREFA:
Escreva o texto acadêmico solicitado seguindo ABNT E especifique a formatação estrutural completa.

REGRAS ABNT:
- Fonte: Arial ou Times New Roman 12pt
- Espaçamento: 1.5 entre linhas
- Alinhamento: Justificado
- Recuo primeira linha: 1.25cm
- Texto acadêmico formal
- Citações quando necessário

IMPORTANTE:
- Escreva de forma acadêmica e formal
- Use vocabulário técnico apropriado
- Mantenha coesão com o texto existente
- Gere texto de qualidade (mínimo 150 palavras)

Retorne APENAS um JSON válido no seguinte formato:
```json
{{
  "content": "O texto completo aqui com múltiplos parágrafos separados por \\n\\n",
  "structure": {{
    "type": "body",
    "formatting": {{
      "font": "Arial",
      "size": 12,
      "bold": false,
      "italic": false,
      "alignment": "justify",
      "spacing": 1.5,
      "indent": 1.25
    }},
    "paragraphs_count": 3,
    "section": "{section_type}"
  }}
}}
```

ATENÇÃO: Retorne APENAS o JSON, sem markdown, sem explicações adicionais.
"""

    try:
        response = model.generate_content(prompt)
        response_text = response.text.strip()

        # Remover markdown code blocks se houver
        if response_text.startswith("```json"):
            response_text = response_text.replace("```json", "").replace("```", "").strip()
        elif response_text.startswith("```"):
            response_text = response_text.replace("```", "").strip()

        result = json.loads(response_text)

        # Validar estrutura
        if "content" not in result or "structure" not in result:
            raise ValueError("Resposta da IA não contém campos obrigatórios")

        return result

    except json.JSONDecodeError as e:
        print(f"Erro ao parsear JSON da IA: {e}")
        print(f"Resposta recebida: {response_text[:500]}")
        # Fallback: retornar texto simples sem estrutura
        return {
            "content": response.text,
            "structure": {
                "type": "body",
                "formatting": {
                    "font": "Arial",
                    "size": 12,
                    "alignment": "justify",
                    "spacing": 1.5,
                    "indent": 1.25
                },
                "paragraphs_count": 1,
                "section": section_type
            }
        }
    except Exception as e:
        print(f"Erro ao gerar texto estruturado: {e}")
        raise


def execute_write_with_structure(
    doc_path: str,
    output_path: str,
    write_result: Dict[str, Any],
    position: str = "fim"
) -> Dict[str, Any]:
    """
    Executa a escrita estruturada no documento

    Args:
        doc_path: Caminho do documento original
        output_path: Caminho de saída
        write_result: Resultado do write_with_structure()
        position: Onde inserir ("fim", "introducao", etc.)

    Returns:
        dict: Resultado da execução
    """
    doc = Document(doc_path)

    content = write_result["content"]
    structure = write_result["structure"]
    formatting = structure.get("formatting", {})

    # Dividir conteúdo em parágrafos
    paragraphs_text = content.split("\n\n")

    # Inserir parágrafos no documento
    inserted_paragraphs = []

    for para_text in paragraphs_text:
        if not para_text.strip():
            continue

        # Adicionar parágrafo
        paragraph = doc.add_paragraph()

        # Adicionar texto
        run = paragraph.add_run(para_text.strip())

        # Aplicar formatação de fonte
        font_name = formatting.get("font", "Arial")
        font_size = formatting.get("size", 12)
        bold = formatting.get("bold", False)
        italic = formatting.get("italic", False)

        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic

        # Aplicar formatação de parágrafo
        alignment = formatting.get("alignment", "justify")
        spacing = formatting.get("spacing", 1.5)
        indent = formatting.get("indent", 1.25)

        # Alinhamento
        alignment_map = {
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT
        }
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

        # Espaçamento entre linhas
        paragraph.paragraph_format.line_spacing = spacing

        # Recuo primeira linha
        if indent > 0:
            paragraph.paragraph_format.first_line_indent = Cm(indent)

        inserted_paragraphs.append({
            "text": para_text.strip()[:100] + "...",
            "formatting_applied": formatting
        })

    # Salvar documento
    doc.save(output_path)

    return {
        "success": True,
        "paragraphs_inserted": len(inserted_paragraphs),
        "total_words": len(content.split()),
        "total_characters": len(content),
        "formatting_applied": formatting,
        "inserted_paragraphs": inserted_paragraphs
    }


def write_structured_streaming(
    document_context: str,
    instruction: str,
    section_type: str,
    document_structure: Dict[str, Any]
):
    """
    Versão streaming da escrita estruturada (para SSE)

    Args:
        document_context: Texto do documento atual
        instruction: Instrução do usuário
        section_type: Tipo da seção
        document_structure: Estrutura completa

    Yields:
        str: Chunks de texto conforme são gerados
    """
    # Verificar se API key está configurada
    if not GOOGLE_API_KEY:
        raise ValueError("GOOGLE_API_KEY não configurada. Configure a variável de ambiente.")

    model = genai.GenerativeModel("gemini-2.5-flash")

    structure_summary = {
        "total_paragraphs": len(document_structure.get("paragraphs", [])),
        "hierarchy": document_structure.get("hierarchy", [])[:5]
    }

    prompt = f"""
Você é um escritor acadêmico especialista em ABNT.

CONTEXTO:
{document_context[-1500:]}

ESTRUTURA:
Total de parágrafos: {structure_summary['total_paragraphs']}

INSTRUÇÃO: {instruction}
SEÇÃO: {section_type}

Escreva texto acadêmico seguindo ABNT:
- Arial/Times 12pt
- Espaçamento 1.5
- Alinhamento justificado
- Recuo 1.25cm
- Mínimo 150 palavras
- Tom formal e técnico

Escreva APENAS o texto, sem JSON, sem formatação extra.
"""

    try:
        response = model.generate_content(prompt, stream=True)

        for chunk in response:
            if chunk.text:
                yield chunk.text

    except Exception as e:
        print(f"Erro no streaming: {e}")
        yield f"[ERRO: {str(e)}]"


def create_action_plan_for_writing(
    content: str,
    formatting: Dict[str, Any],
    position: str = "fim"
) -> List[Dict[str, Any]]:
    """
    Cria plano de ação para inserção de texto estruturado

    Args:
        content: Conteúdo a ser inserido
        formatting: Formatação a ser aplicada
        position: Posição de inserção

    Returns:
        list: Lista de ações executáveis
    """
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

    action_plan = []

    for i, para_text in enumerate(paragraphs):
        action_plan.append({
            "action": "insert_paragraph",
            "priority": i + 1,
            "target": position,
            "params": {
                "text": para_text,
                "font": formatting.get("font", "Arial"),
                "size": formatting.get("size", 12),
                "bold": formatting.get("bold", False),
                "italic": formatting.get("italic", False),
                "alignment": formatting.get("alignment", "justify"),
                "spacing": formatting.get("spacing", 1.5),
                "indent": formatting.get("indent", 1.25)
            },
            "description": f"Inserir parágrafo {i+1} com formatação ABNT"
        })

    return action_plan
