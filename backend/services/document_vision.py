"""
NORMAEX 2.0 - Document Vision Engine
Extração completa da estrutura do documento (DOCX + PDF)
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
import os
from typing import Dict, List, Optional, Any


def extract_complete_structure(docx_path: str) -> Dict[str, Any]:
    """
    Extrai TODA a estrutura do documento em JSON

    Args:
        docx_path: Caminho para o arquivo DOCX

    Returns:
        dict: Estrutura completa do documento
    """
    doc = Document(docx_path)

    structure = {
        "metadata": extract_metadata(doc),
        "sections": extract_sections(doc),
        "paragraphs": extract_paragraphs(doc),
        "styles": extract_all_styles(doc),
        "hierarchy": detect_hierarchy(doc),
        "statistics": calculate_statistics(doc)
    }

    return structure


def extract_metadata(doc: Document) -> Dict[str, Any]:
    """
    Extrai metadados do documento
    """
    core_props = doc.core_properties

    metadata = {
        "title": core_props.title or "Sem título",
        "author": core_props.author or "Desconhecido",
        "subject": core_props.subject or "",
        "created": str(core_props.created) if core_props.created else None,
        "modified": str(core_props.modified) if core_props.modified else None,
        "last_modified_by": core_props.last_modified_by or "Desconhecido"
    }

    return metadata


def extract_sections(doc: Document) -> List[Dict[str, Any]]:
    """
    Extrai informações de seções (margens, orientação, tamanho)
    """
    sections = []

    for idx, section in enumerate(doc.sections):
        section_info = {
            "index": idx,
            "margins": {
                "top": round(section.top_margin.cm, 2) if section.top_margin else None,
                "bottom": round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
                "left": round(section.left_margin.cm, 2) if section.left_margin else None,
                "right": round(section.right_margin.cm, 2) if section.right_margin else None,
            },
            "page_size": {
                "width": round(section.page_width.cm, 2) if section.page_width else None,
                "height": round(section.page_height.cm, 2) if section.page_height else None,
            },
            "orientation": str(section.orientation),
        }
        sections.append(section_info)

    return sections


def extract_paragraphs(doc: Document) -> List[Dict[str, Any]]:
    """
    Extrai cada parágrafo com TODOS os detalhes
    """
    paragraphs = []

    for idx, para in enumerate(doc.paragraphs):
        # Extrair formatação do parágrafo
        para_format = para.paragraph_format

        para_info = {
            "index": idx,
            "text": para.text,
            "length": len(para.text),
            "style": {
                "name": para.style.name,
                "font": extract_font_info(para.style.font) if hasattr(para.style, 'font') else {},
                "alignment": str(para.alignment) if para.alignment else "None",
                "spacing": {
                    "before": round(para_format.space_before.pt, 2) if para_format.space_before else None,
                    "after": round(para_format.space_after.pt, 2) if para_format.space_after else None,
                    "line_spacing": para_format.line_spacing if para_format.line_spacing else None,
                    "line_spacing_rule": str(para_format.line_spacing_rule) if para_format.line_spacing_rule else None,
                },
                "indent": {
                    "left": round(para_format.left_indent.cm, 2) if para_format.left_indent else None,
                    "right": round(para_format.right_indent.cm, 2) if para_format.right_indent else None,
                    "first_line": round(para_format.first_line_indent.cm, 2) if para_format.first_line_indent else None,
                }
            },
            "runs": []
        }

        # Extrair runs (trechos de texto com formatação específica)
        for run_idx, run in enumerate(para.runs):
            run_info = {
                "index": run_idx,
                "text": run.text,
                "length": len(run.text),
                "bold": run.bold if run.bold is not None else False,
                "italic": run.italic if run.italic is not None else False,
                "underline": run.underline if run.underline is not None else False,
                "font": extract_font_info(run.font)
            }
            para_info["runs"].append(run_info)

        paragraphs.append(para_info)

    return paragraphs


def extract_font_info(font) -> Dict[str, Any]:
    """
    Extrai informações de fonte
    """
    return {
        "name": font.name,
        "size": round(font.size.pt, 2) if font.size else None,
        "bold": font.bold if font.bold is not None else False,
        "italic": font.italic if font.italic is not None else False,
        "color": str(font.color.rgb) if hasattr(font.color, 'rgb') and font.color.rgb else None,
    }


def extract_all_styles(doc: Document) -> Dict[str, Any]:
    """
    Extrai todos os estilos disponíveis no documento
    """
    styles_info = {
        "total_styles": len(doc.styles),
        "paragraph_styles": [],
        "character_styles": [],
    }

    for style in doc.styles:
        style_data = {
            "name": style.name,
            "type": str(style.type),
            "builtin": style.builtin,
        }

        # Classificar por tipo
        if "PARAGRAPH" in str(style.type):
            styles_info["paragraph_styles"].append(style_data)
        elif "CHARACTER" in str(style.type):
            styles_info["character_styles"].append(style_data)

    return styles_info


def detect_hierarchy(doc: Document) -> List[Dict[str, Any]]:
    """
    Detecta hierarquia do documento (títulos, subtítulos)
    """
    hierarchy = []

    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name.lower()

        # Detectar headings
        if 'heading' in style_name or 'título' in style_name:
            level = 1

            # Extrair nível do heading
            if 'heading 1' in style_name or 'título 1' in style_name:
                level = 1
            elif 'heading 2' in style_name or 'título 2' in style_name:
                level = 2
            elif 'heading 3' in style_name or 'título 3' in style_name:
                level = 3
            elif 'heading' in style_name:
                try:
                    level = int(style_name.split()[-1])
                except:
                    level = 1

            hierarchy.append({
                "paragraph_index": idx,
                "level": level,
                "text": para.text,
                "style": para.style.name,
                "length": len(para.text)
            })

    return hierarchy


def calculate_statistics(doc: Document) -> Dict[str, Any]:
    """
    Calcula estatísticas do documento
    """
    total_chars = 0
    total_words = 0
    total_paragraphs = len(doc.paragraphs)
    non_empty_paragraphs = 0

    fonts_used = {}
    font_sizes_used = {}

    for para in doc.paragraphs:
        if para.text.strip():
            non_empty_paragraphs += 1
            total_chars += len(para.text)
            total_words += len(para.text.split())

        # Contar fontes usadas
        for run in para.runs:
            if run.font.name:
                fonts_used[run.font.name] = fonts_used.get(run.font.name, 0) + 1

            if run.font.size:
                size = round(run.font.size.pt, 1)
                font_sizes_used[size] = font_sizes_used.get(size, 0) + 1

    return {
        "total_paragraphs": total_paragraphs,
        "non_empty_paragraphs": non_empty_paragraphs,
        "total_characters": total_chars,
        "total_words": total_words,
        "avg_words_per_paragraph": round(total_words / non_empty_paragraphs, 2) if non_empty_paragraphs > 0 else 0,
        "fonts_used": fonts_used,
        "font_sizes_used": font_sizes_used
    }


def extract_visual_layout(pdf_path: str) -> Dict[str, Any]:
    """
    Usa PyMuPDF para extrair coordenadas reais e layout visual do PDF

    Args:
        pdf_path: Caminho para o arquivo PDF

    Returns:
        dict: Dados visuais do PDF com coordenadas
    """
    doc = fitz.open(pdf_path)

    visual_data = {
        "total_pages": len(doc),
        "pages": []
    }

    for page_num in range(len(doc)):
        page = doc[page_num]

        page_data = {
            "page_number": page_num + 1,
            "size": {
                "width": round(page.rect.width, 2),
                "height": round(page.rect.height, 2)
            },
            "text_blocks": [],
            "images": []
        }

        # Extrair blocos de texto com coordenadas
        blocks = page.get_text("dict")["blocks"]

        for block in blocks:
            if "lines" in block:  # Bloco de texto
                for line in block["lines"]:
                    for span in line["spans"]:
                        text_info = {
                            "text": span["text"],
                            "bbox": {
                                "x0": round(span["bbox"][0], 2),
                                "y0": round(span["bbox"][1], 2),
                                "x1": round(span["bbox"][2], 2),
                                "y1": round(span["bbox"][3], 2)
                            },
                            "font": span["font"],
                            "size": round(span["size"], 2),
                            "color": span["color"],
                            "flags": span["flags"]  # Bold, italic, etc.
                        }
                        page_data["text_blocks"].append(text_info)

            elif "image" in block:  # Bloco de imagem
                image_info = {
                    "bbox": {
                        "x0": round(block["bbox"][0], 2),
                        "y0": round(block["bbox"][1], 2),
                        "x1": round(block["bbox"][2], 2),
                        "y1": round(block["bbox"][3], 2)
                    },
                    "width": round(block["width"], 2),
                    "height": round(block["height"], 2)
                }
                page_data["images"].append(image_info)

        visual_data["pages"].append(page_data)

    doc.close()
    return visual_data


def calculate_visual_margins(visual_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Calcula margens reais baseado nas coordenadas visuais do PDF
    """
    if not visual_data["pages"]:
        return {}

    first_page = visual_data["pages"][0]

    if not first_page["text_blocks"]:
        return {}

    # Pegar coordenadas do primeiro bloco de texto
    text_blocks = first_page["text_blocks"]

    # Calcular margens (72 points = 1 inch = 2.54 cm)
    left_positions = [block["bbox"]["x0"] for block in text_blocks]
    top_positions = [block["bbox"]["y0"] for block in text_blocks]
    right_positions = [block["bbox"]["x1"] for block in text_blocks]

    page_width = first_page["size"]["width"]
    page_height = first_page["size"]["height"]

    # Margens em cm
    left_margin_cm = (min(left_positions) / 72) * 2.54 if left_positions else 0
    top_margin_cm = (min(top_positions) / 72) * 2.54 if top_positions else 0
    right_margin_cm = ((page_width - max(right_positions)) / 72) * 2.54 if right_positions else 0

    return {
        "left": round(left_margin_cm, 2),
        "top": round(top_margin_cm, 2),
        "right": round(right_margin_cm, 2),
        "measured_from": "pdf_coordinates"
    }


def merge_docx_and_pdf_data(docx_structure: Dict[str, Any], pdf_visual: Dict[str, Any]) -> Dict[str, Any]:
    """
    Combina dados estruturais do DOCX com dados visuais do PDF

    Args:
        docx_structure: Estrutura extraída do DOCX
        pdf_visual: Dados visuais extraídos do PDF

    Returns:
        dict: Visão completa combinada
    """
    complete_vision = {
        "structure": docx_structure,
        "visual": pdf_visual,
        "visual_margins": calculate_visual_margins(pdf_visual),
        "analysis": {
            "total_elements": {
                "paragraphs": len(docx_structure["paragraphs"]),
                "sections": len(docx_structure["sections"]),
                "hierarchy_levels": len(docx_structure["hierarchy"]),
                "pages": pdf_visual["total_pages"]
            },
            "document_type": detect_document_type(docx_structure),
            "abnt_compliance": quick_abnt_check(docx_structure)
        }
    }

    return complete_vision


def detect_document_type(structure: Dict[str, Any]) -> str:
    """
    Detecta o tipo de documento baseado na estrutura
    """
    hierarchy = structure["hierarchy"]
    total_words = structure["statistics"]["total_words"]

    if total_words < 500:
        return "documento_curto"
    elif total_words < 3000:
        return "artigo_ou_relatorio"
    elif total_words < 10000:
        return "monografia_ou_tcc"
    else:
        return "dissertacao_ou_tese"


def quick_abnt_check(structure: Dict[str, Any]) -> Dict[str, Any]:
    """
    Verificação rápida de conformidade ABNT
    """
    issues = []

    # Verificar margens
    if structure["sections"]:
        margins = structure["sections"][0]["margins"]

        if margins["top"] != 3.0:
            issues.append(f"Margem superior: {margins['top']}cm (deve ser 3cm)")
        if margins["left"] != 3.0:
            issues.append(f"Margem esquerda: {margins['left']}cm (deve ser 3cm)")
        if margins["bottom"] != 2.0:
            issues.append(f"Margem inferior: {margins['bottom']}cm (deve ser 2cm)")
        if margins["right"] != 2.0:
            issues.append(f"Margem direita: {margins['right']}cm (deve ser 2cm)")

    # Verificar fontes
    fonts = structure["statistics"]["fonts_used"]
    if fonts:
        main_font = max(fonts, key=fonts.get)
        if main_font not in ["Arial", "Times New Roman"]:
            issues.append(f"Fonte principal '{main_font}' não é Arial ou Times New Roman")

    # Verificar tamanho de fonte
    sizes = structure["statistics"]["font_sizes_used"]
    if sizes:
        main_size = max(sizes, key=sizes.get)
        if main_size != 12.0:
            issues.append(f"Tamanho de fonte principal {main_size}pt (deve ser 12pt)")

    return {
        "compliant": len(issues) == 0,
        "total_issues": len(issues),
        "issues": issues
    }


def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """
    Converte DOCX para PDF

    Args:
        docx_path: Caminho do arquivo DOCX
        pdf_path: Caminho de saída do PDF

    Returns:
        bool: True se conversão bem-sucedida
    """
    # Verificar se arquivo de entrada existe
    if not os.path.exists(docx_path):
        return False

    try:
        # Tentar usar docx2pdf (Windows)
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            # Verificar se PDF foi criado
            return os.path.exists(pdf_path)
        except ImportError:
            pass  # docx2pdf não disponível
        except Exception:
            # Se ocorreu erro mas PDF foi criado, considerar sucesso
            if os.path.exists(pdf_path):
                return True
            # Caso contrário, tentar LibreOffice

        # Tentar usar LibreOffice (Linux/Mac)
        import subprocess
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
             os.path.dirname(pdf_path), docx_path],
            capture_output=True,
            timeout=30
        )
        return result.returncode == 0 and os.path.exists(pdf_path)

    except Exception:
        # Em último caso, verificar se o PDF existe
        return os.path.exists(pdf_path)
