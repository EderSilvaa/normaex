from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .ai import organize_references_ai
import re


# Padrões para identificar seções comuns em TCCs
SECTION_PATTERNS = {
    "introducao": [r"introdu[çc][ãa]o", r"1\.?\s*introdu"],
    "referencial": [r"referencial\s*te[óo]rico", r"revis[ãa]o\s*(da\s*)?literatura", r"fundamenta[çc][ãa]o"],
    "metodologia": [r"metodologia", r"m[ée]todos?", r"procedimentos?\s*metodol[óo]gicos?"],
    "resultados": [r"resultados?", r"an[áa]lise\s*(dos\s*)?resultados?", r"discuss[ãa]o"],
    "conclusao": [r"conclus[ãa]o", r"considera[çc][õo]es\s*finais"],
    "referencias": [r"refer[êe]ncias?", r"bibliografia"]
}


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def analyze_document(file_path: str) -> dict:
    """
    Analisa o documento usando FASE 2 (AI Structural) e retorna as alterações necessárias SEM aplicar.
    Integrado com sistema de análise inteligente do Normaex 2.0.
    """
    from services.document_vision import extract_complete_structure
    from services.ai_structural import analyze_document_with_ai, detect_style_inconsistencies

    doc = Document(file_path)
    issues = []

    # === FASE 2: ANÁLISE ESTRUTURAL COM IA ===
    try:
        # Extrair estrutura completa do documento
        complete_vision = extract_complete_structure(file_path)

        # Análise inteligente com IA
        ai_analysis = analyze_document_with_ai(complete_vision)

        # Converter análise da IA para formato de issues
        for ai_issue in ai_analysis.get("ai_analysis", {}).get("abnt_issues", []):
            issues.append({
                "id": f"ai_{ai_issue['category']}",
                "category": ai_issue['category'].capitalize(),
                "description": ai_issue['description'],
                "severity": ai_issue['severity']
            })

        # Adicionar inconsistências detectadas
        for inconsistency in ai_analysis.get("inconsistencies", []):
            issues.append({
                "id": f"inconsistency_{inconsistency['type']}",
                "category": inconsistency['type'].replace('_', ' ').title(),
                "description": inconsistency['description'],
                "recommendation": inconsistency['recommendation'],
                "affected_count": inconsistency.get('affected_count', 0),
                "severity": inconsistency['severity']
            })

    except Exception as e:
        print(f"Erro na análise com IA (Fase 2): {e}")
        # Fallback para análise básica se IA falhar

    # === ANÁLISE BÁSICA (SEMPRE EXECUTADA) ===
    # Verificar margens (SEM DUPLICAR)
    section = doc.sections[0]  # Apenas primeira seção
    margin_issues = []

    if section.top_margin != Cm(3):
        margin_issues.append(f"Superior: {section.top_margin.cm:.1f}cm (esperado: 3cm)")
    if section.left_margin != Cm(3):
        margin_issues.append(f"Esquerda: {section.left_margin.cm:.1f}cm (esperado: 3cm)")
    if section.bottom_margin != Cm(2):
        margin_issues.append(f"Inferior: {section.bottom_margin.cm:.1f}cm (esperado: 2cm)")
    if section.right_margin != Cm(2):
        margin_issues.append(f"Direita: {section.right_margin.cm:.1f}cm (esperado: 2cm)")

    if margin_issues:
        issues.append({
            "id": "margins",
            "category": "Margens",
            "description": "Margens fora do padrão ABNT: " + ", ".join(margin_issues),
            "severity": "high"
        })

    # Verificar fonte padrão
    style = doc.styles['Normal']
    font = style.font

    if font.name and font.name != 'Arial':
        issues.append({
            "id": "font_family",
            "category": "Fonte",
            "description": f"Fonte padrão '{font.name}' não é Arial",
            "expected": "Arial",
            "severity": "medium"
        })

    if font.size and font.size != Pt(12):
        issues.append({
            "id": "font_size",
            "category": "Fonte",
            "description": f"Tamanho da fonte {font.size.pt:.0f}pt não é 12pt",
            "expected": "12pt",
            "severity": "medium"
        })

    # Contar parágrafos com problemas E REGISTRAR QUAIS
    wrong_style_paragraphs = []
    wrong_runs_details = []
    no_indent_paragraphs = []

    for idx, paragraph in enumerate(doc.paragraphs):
        para_num = idx + 1
        para_preview = paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text

        # Verificar estilo
        if paragraph.style.name != 'Normal' and paragraph.text.strip():
            if not paragraph.style.name.lower().startswith('heading'):
                wrong_style_paragraphs.append({
                    "paragraph": para_num,
                    "style": paragraph.style.name,
                    "preview": para_preview
                })

        # Verificar recuo
        if paragraph.paragraph_format.first_line_indent != Cm(1.25):
            if not paragraph.style.name.lower().startswith('heading') and paragraph.text.strip():
                current_indent = paragraph.paragraph_format.first_line_indent.cm if paragraph.paragraph_format.first_line_indent else 0
                no_indent_paragraphs.append({
                    "paragraph": para_num,
                    "current_indent": f"{current_indent:.2f}cm",
                    "preview": para_preview
                })

        # Verificar formatação manual nos runs
        for run_idx, run in enumerate(paragraph.runs):
            if run.font.name and run.font.name != 'Arial':
                wrong_runs_details.append({
                    "paragraph": para_num,
                    "run": run_idx + 1,
                    "issue": f"Fonte '{run.font.name}' ao invés de Arial",
                    "text_preview": run.text[:30] + "..." if len(run.text) > 30 else run.text
                })
            elif run.font.size and run.font.size != Pt(12):
                wrong_runs_details.append({
                    "paragraph": para_num,
                    "run": run_idx + 1,
                    "issue": f"Tamanho {run.font.size.pt:.0f}pt ao invés de 12pt",
                    "text_preview": run.text[:30] + "..." if len(run.text) > 30 else run.text
                })

    # Adicionar issues com DETALHES dos parágrafos afetados
    if wrong_style_paragraphs:
        # Mostrar primeiros 5 parágrafos como exemplo
        examples = wrong_style_paragraphs[:5]
        examples_text = ", ".join([f"§{p['paragraph']}" for p in examples])
        more_text = f" e mais {len(wrong_style_paragraphs) - 5}" if len(wrong_style_paragraphs) > 5 else ""

        issues.append({
            "id": "paragraph_styles",
            "category": "Estilos",
            "description": f"{len(wrong_style_paragraphs)} parágrafos usando estilos não-padrão (ex: {examples_text}{more_text})",
            "affected_count": len(wrong_style_paragraphs),
            "affected_paragraphs": wrong_style_paragraphs,
            "severity": "medium"
        })

    if wrong_runs_details:
        # Mostrar primeiros 5 exemplos
        examples = wrong_runs_details[:5]
        examples_text = ", ".join([f"§{d['paragraph']}" for d in examples])
        more_text = f" e mais {len(wrong_runs_details) - 5}" if len(wrong_runs_details) > 5 else ""

        issues.append({
            "id": "manual_formatting",
            "category": "Formatação Manual",
            "description": f"{len(wrong_runs_details)} trechos com formatação manual inconsistente (ex: {examples_text}{more_text})",
            "affected_count": len(wrong_runs_details),
            "affected_details": wrong_runs_details[:20],  # Primeiros 20 para não sobrecarregar
            "severity": "low"
        })

    if no_indent_paragraphs:
        # Mostrar primeiros 5 parágrafos
        examples = no_indent_paragraphs[:5]
        examples_text = ", ".join([f"§{p['paragraph']}" for p in examples])
        more_text = f" e mais {len(no_indent_paragraphs) - 5}" if len(no_indent_paragraphs) > 5 else ""

        issues.append({
            "id": "no_indent",
            "category": "Recuo",
            "description": f"{len(no_indent_paragraphs)} parágrafos sem recuo de 1.25cm (ex: {examples_text}{more_text})",
            "affected_count": len(no_indent_paragraphs),
            "affected_paragraphs": no_indent_paragraphs[:20],  # Primeiros 20
            "severity": "medium"
        })

    # Verificar referências
    has_references = False
    ref_count = 0
    for i, p in enumerate(doc.paragraphs):
        if p.text.lower().strip() in ['referências', 'referencias', 'bibliografia']:
            has_references = True
            for j in range(i + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():
                    ref_count += 1
            break

    if has_references and ref_count > 0:
        issues.append({
            "id": "references",
            "category": "Referências",
            "description": f"Seção de referências com {ref_count} itens (verificar ordem alfabética)",
            "affected_count": ref_count,
            "severity": "low"
        })

    # Resumo
    summary = {
        "total_issues": len(issues),
        "categories": list(set(issue["category"] for issue in issues)),
        "issues": issues,
        "analysis_method": "Normaex 2.0 - Fase 2 (AI Structural Analysis)"
    }

    return summary


def format_abnt(file_path: str, output_path: str):
    """
    Aplica todas as formatações ABNT no documento.
    """
    doc = Document(file_path)

    changes = []

    # Configurar margens
    sections = doc.sections
    for i, section in enumerate(sections):
        if section.top_margin != Cm(3) or section.left_margin != Cm(3):
             changes.append(f"Seção {i+1}: Margens superior/esquerda ajustadas para 3cm.")
        if section.bottom_margin != Cm(2) or section.right_margin != Cm(2):
             changes.append(f"Seção {i+1}: Margens inferior/direita ajustadas para 2cm.")

        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Configurar estilo normal (parágrafos)
    style = doc.styles['Normal']
    font = style.font

    if font.name != 'Arial':
        changes.append("Fonte padrão alterada para Arial.")
    font.name = 'Arial'

    if font.size != Pt(12):
        changes.append("Tamanho da fonte ajustado para 12pt.")
    font.size = Pt(12)

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.first_line_indent = Cm(1.25)

    changes.append("Espaçamento entre linhas definido para 1.5 e recuo de 1.25cm aplicado.")

    # Iterar parágrafos para garantir formatação
    count_fixed = 0
    count_runs_fixed = 0

    for paragraph in doc.paragraphs:
        if paragraph.style.name != 'Normal':
            paragraph.style = doc.styles['Normal']
            count_fixed += 1

        for run in paragraph.runs:
            if run.font.name != 'Arial' or run.font.size != Pt(12):
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                count_runs_fixed += 1

    if count_fixed > 0:
        changes.append(f"{count_fixed} parágrafos com estilo incorreto foram normalizados.")

    if count_runs_fixed > 0:
        changes.append(f"{count_runs_fixed} trechos de texto com formatação manual foram corrigidos.")

    # Organizar referências
    ref_start_index = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.lower().strip() in ['referências', 'referencias', 'bibliografia']:
            ref_start_index = i
            break

    if ref_start_index != -1:
        ref_texts = []
        for i in range(ref_start_index + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip():
                ref_texts.append(doc.paragraphs[i].text)

        if ref_texts:
            full_ref_text = "\n".join(ref_texts)
            organized_text = organize_references_ai(full_ref_text)

            for i in range(len(doc.paragraphs) - 1, ref_start_index, -1):
                p = doc.paragraphs[i]
                p._element.getparent().remove(p._element)

            for line in organized_text.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.style = doc.styles['Normal']

            changes.append("Referências bibliográficas organizadas e formatadas pela IA.")

    doc.save(output_path)
    return output_path, changes


def identify_sections(file_path: str) -> list:
    """
    Identifica as seções do documento e suas posições.
    """
    doc = Document(file_path)
    sections = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if not text:
            continue

        for section_name, patterns in SECTION_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    sections.append({
                        "name": section_name,
                        "title": para.text.strip(),
                        "index": i
                    })
                    break

    return sections


def insert_text_after_section(file_path: str, output_path: str, section_name: str, text_to_insert: str) -> tuple:
    """
    Insere texto após uma seção específica do documento.
    """
    doc = Document(file_path)
    sections = identify_sections(file_path)

    # Encontrar a seção alvo
    target_section = None
    next_section_index = len(doc.paragraphs)

    for i, section in enumerate(sections):
        if section["name"] == section_name.lower():
            target_section = section
            # Encontrar o índice da próxima seção
            if i + 1 < len(sections):
                next_section_index = sections[i + 1]["index"]
            break

    if not target_section:
        # Se não encontrar a seção, adiciona no final antes das referências
        ref_index = len(doc.paragraphs)
        for i, para in enumerate(doc.paragraphs):
            if re.search(r"refer[êe]ncias?|bibliografia", para.text.lower()):
                ref_index = i
                break
        insert_index = ref_index
    else:
        # Inserir antes da próxima seção
        insert_index = next_section_index

    # Criar parágrafos para o novo texto
    paragraphs_to_insert = text_to_insert.split('\n\n')

    # Inserir os parágrafos
    for j, para_text in enumerate(paragraphs_to_insert):
        if para_text.strip():
            # Criar novo parágrafo
            new_para = doc.add_paragraph()
            new_para.text = para_text.strip()
            new_para.style = doc.styles['Normal']

            # Aplicar formatação ABNT
            new_para.paragraph_format.line_spacing = 1.5
            new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            new_para.paragraph_format.first_line_indent = Cm(1.25)

            for run in new_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)

            # Mover para a posição correta
            target_element = doc.paragraphs[insert_index + j]._element
            new_para._element.addprevious(target_element)

    doc.save(output_path)
    return output_path, f"Texto inserido após a seção '{section_name}'"


def insert_text_at_end(file_path: str, output_path: str, text_to_insert: str) -> tuple:
    """
    Insere texto no final do documento (antes das referências se existirem).
    """
    doc = Document(file_path)

    # Encontrar onde inserir (antes de referências)
    insert_index = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs):
        if re.search(r"refer[êe]ncias?|bibliografia", para.text.lower()):
            insert_index = i
            break

    # Criar parágrafos para o novo texto
    paragraphs_to_insert = text_to_insert.split('\n\n')

    for para_text in paragraphs_to_insert:
        if para_text.strip():
            new_para = doc.add_paragraph()
            new_para.text = para_text.strip()
            new_para.style = doc.styles['Normal']

            # Aplicar formatação ABNT
            new_para.paragraph_format.line_spacing = 1.5
            new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            new_para.paragraph_format.first_line_indent = Cm(1.25)

            for run in new_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)

    doc.save(output_path)
    return output_path, "Texto inserido no documento"


def get_document_structure(file_path: str) -> dict:
    """
    Retorna a estrutura do documento com as seções identificadas.
    """
    doc = Document(file_path)
    sections = identify_sections(file_path)

    # Contar parágrafos por seção
    structure = {
        "total_paragraphs": len(doc.paragraphs),
        "sections": []
    }

    for i, section in enumerate(sections):
        next_index = sections[i + 1]["index"] if i + 1 < len(sections) else len(doc.paragraphs)
        paragraph_count = next_index - section["index"] - 1

        structure["sections"].append({
            "name": section["name"],
            "title": section["title"],
            "paragraph_count": paragraph_count
        })

    return structure
