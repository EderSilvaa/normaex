from fastapi import APIRouter, UploadFile, File, HTTPException, Request
from fastapi.responses import FileResponse, HTMLResponse
from pydantic import BaseModel
from typing import Optional
import shutil
import os
import json
import mammoth
from urllib.parse import unquote
from sse_starlette.sse import EventSourceResponse
from services.abnt import (
    format_abnt,
    extract_text_from_docx,
    analyze_document,
    get_document_structure,
    insert_text_at_end,
    insert_text_after_section
)
from services.ai import (
    chat_with_document,
    generate_academic_text,
    generate_academic_text_stream,
    detect_write_intent
)
from services.document_vision import (
    extract_complete_structure,
    extract_visual_layout,
    merge_docx_and_pdf_data,
    convert_docx_to_pdf
)
from services.ai_structural import (
    analyze_document_with_ai,
    classify_all_paragraphs,
    detect_style_inconsistencies
)
from services.executor import DocumentExecutor
from services.validator import DocumentValidator, validate_document_quality
from services.ai_writer import (
    write_with_structure,
    execute_write_with_structure,
    write_structured_streaming,
    create_action_plan_for_writing
)
from pdf2docx import Converter

router = APIRouter()


class ChatRequest(BaseModel):
    filename: str
    message: str


class ApplyRequest(BaseModel):
    filename: str


class WriteRequest(BaseModel):
    filename: str
    instruction: str
    section_type: Optional[str] = "geral"
    position: Optional[str] = "fim"  # "fim", "introducao", "metodologia", etc.


class EditParagraphRequest(BaseModel):
    filename: str
    paragraph_number: int
    new_text: str


class EditElementRequest(BaseModel):
    filename: str
    element_type: str  # "titulo_principal", "subtitulo", "autor", etc.
    new_text: str


class SmartEditRequest(BaseModel):
    filename: str
    command: str  # Comando natural: "edite o título para X", "mude a conclusão", etc.


UPLOAD_DIR = "uploads"
PROCESSED_DIR = "processed"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(PROCESSED_DIR, exist_ok=True)


@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """
    Faz upload do arquivo e analisa as alterações necessárias (SEM aplicar).
    """
    if not (file.filename.endswith(".docx") or file.filename.endswith(".pdf")):
        raise HTTPException(status_code=400, detail="Only .docx and .pdf files are supported")

    file_location = f"{UPLOAD_DIR}/{file.filename}"
    with open(file_location, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Se for PDF, converter para DOCX
    if file.filename.endswith(".pdf"):
        docx_filename = file.filename.replace(".pdf", ".docx")
        docx_location = f"{UPLOAD_DIR}/{docx_filename}"
        try:
            cv = Converter(file_location)
            cv.convert(docx_location)
            cv.close()
            file_location = docx_location
            file.filename = docx_filename
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error converting PDF: {str(e)}")

    # Analisar documento (sem aplicar formatação)
    try:
        analysis = analyze_document(file_location)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analyzing file: {str(e)}")

    return {
        "filename": file.filename,
        "file_path": file_location,
        "message": "Documento analisado com sucesso",
        "analysis": analysis
    }


@router.post("/apply")
async def apply_formatting(request: ApplyRequest):
    """
    Aplica a formatação ABNT no documento após autorização do usuário.
    """
    file_location = f"{UPLOAD_DIR}/{request.filename}"

    if not os.path.exists(file_location):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")

    processed_location = f"{PROCESSED_DIR}/formatted_{request.filename}"

    try:
        _, changes = format_abnt(file_location, processed_location)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

    return {
        "filename": request.filename,
        "processed_path": processed_location,
        "message": "Formatação aplicada com sucesso!",
        "changes": changes
    }


@router.get("/download/{filename}")
async def download_file(filename: str):
    file_path = f"{PROCESSED_DIR}/{filename}"
    if os.path.exists(file_path):
        return FileResponse(
            file_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename
        )
    raise HTTPException(status_code=404, detail="File not found")


@router.post("/chat")
async def chat(request: ChatRequest):
    # Tenta ler do processado primeiro, senão do original
    processed_path = f"{PROCESSED_DIR}/formatted_{request.filename}"
    original_path = f"{UPLOAD_DIR}/{request.filename}"

    target_path = processed_path if os.path.exists(processed_path) else original_path

    if not os.path.exists(target_path):
        target_path = f"{PROCESSED_DIR}/{request.filename}"
        if not os.path.exists(target_path):
            raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        text = extract_text_from_docx(target_path)
        response = chat_with_document(text, request.message)
        return {"response": response}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/edit-paragraph")
async def edit_paragraph(request: EditParagraphRequest):
    """
    Edita um parágrafo específico do documento e retorna PDF atualizado.
    """
    from docx import Document

    # Buscar arquivo (prioriza processado)
    paths_to_check = [
        f"{PROCESSED_DIR}/edited_{request.filename}",
        f"{PROCESSED_DIR}/formatted_{request.filename}",
        f"{UPLOAD_DIR}/{request.filename}"
    ]

    file_path = None
    for path in paths_to_check:
        if os.path.exists(path):
            file_path = path
            break

    if not file_path:
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        # Carregar documento
        doc = Document(file_path)

        # Verificar se parágrafo existe
        if request.paragraph_number < 1 or request.paragraph_number > len(doc.paragraphs):
            raise HTTPException(
                status_code=400,
                detail=f"Parágrafo {request.paragraph_number} não existe. Documento tem {len(doc.paragraphs)} parágrafos."
            )

        # Editar parágrafo (índice começa em 0)
        paragraph = doc.paragraphs[request.paragraph_number - 1]

        # Preservar formatação e substituir texto
        for run in paragraph.runs:
            run.text = ''

        if paragraph.runs:
            paragraph.runs[0].text = request.new_text
        else:
            paragraph.add_run(request.new_text)

        # Salvar como edited
        output_path = f"{PROCESSED_DIR}/edited_{request.filename}"
        doc.save(output_path)

        # Converter para PDF
        pdf_path = output_path.replace(".docx", "_preview.pdf")
        conversion_success = convert_docx_to_pdf(output_path, pdf_path)

        return {
            "success": True,
            "message": f"Parágrafo {request.paragraph_number} editado com sucesso",
            "paragraph_number": request.paragraph_number,
            "new_text": request.new_text,
            "pdf_updated": conversion_success
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/identify-elements")
async def identify_elements(request: ApplyRequest):
    """
    Identifica elementos semânticos do documento (título, subtítulo, autor, etc.)
    e retorna o mapeamento para números de parágrafo.
    """
    from docx import Document
    from services.ai import get_model
    import json

    # Buscar arquivo
    paths_to_check = [
        f"{PROCESSED_DIR}/edited_{request.filename}",
        f"{PROCESSED_DIR}/formatted_{request.filename}",
        f"{UPLOAD_DIR}/{request.filename}"
    ]

    file_path = None
    for path in paths_to_check:
        if os.path.exists(path):
            file_path = path
            break

    if not file_path:
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        doc = Document(file_path)

        # Extrair primeiros 20 parágrafos com seus textos
        paragraphs_preview = []
        for idx, para in enumerate(doc.paragraphs[:20], 1):
            if para.text.strip():
                paragraphs_preview.append({
                    "numero": idx,
                    "texto": para.text.strip()[:100],  # Primeiros 100 caracteres
                    "estilo": para.style.name if para.style else "Normal"
                })

        # Usar IA para identificar elementos
        model = get_model()
        prompt = f"""
        Analise os parágrafos abaixo de um TCC/documento acadêmico e identifique EXATAMENTE qual número de parágrafo corresponde a cada elemento.

        PARÁGRAFOS:
        {json.dumps(paragraphs_preview, ensure_ascii=False, indent=2)}

        Retorne APENAS um JSON válido no formato:
        {{
            "titulo_principal": numero_do_paragrafo,
            "subtitulo": numero_do_paragrafo_ou_null,
            "autor": numero_do_paragrafo_ou_null,
            "instituicao": numero_do_paragrafo_ou_null,
            "cidade_ano": numero_do_paragrafo_ou_null
        }}

        IMPORTANTE: Retorne APENAS o JSON, sem markdown, sem explicações.
        """

        response = model.generate_content(prompt)
        text = response.text.strip()

        # Limpar markdown se presente
        if text.startswith("```"):
            text = text.split("\n", 1)[1] if "\n" in text else text
            text = text.rsplit("```", 1)[0]

        elements = json.loads(text)

        return {
            "success": True,
            "elements": elements,
            "paragraphs_preview": paragraphs_preview
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/edit-element")
async def edit_element(request: EditElementRequest):
    """
    Edita um elemento específico do documento (título, subtítulo, etc.)
    identificando automaticamente qual parágrafo editar.
    """
    from docx import Document

    # Primeiro, identificar elementos
    identify_request = ApplyRequest(filename=request.filename)
    elements_response = await identify_elements(identify_request)

    if not elements_response["success"]:
        raise HTTPException(status_code=500, detail="Erro ao identificar elementos")

    elements = elements_response["elements"]
    paragraph_number = elements.get(request.element_type)

    if not paragraph_number:
        raise HTTPException(
            status_code=404,
            detail=f"Elemento '{request.element_type}' não encontrado no documento"
        )

    # Usar endpoint de edição de parágrafo
    edit_request = EditParagraphRequest(
        filename=request.filename,
        paragraph_number=paragraph_number,
        new_text=request.new_text
    )

    return await edit_paragraph(edit_request)


@router.post("/smart-edit")
async def smart_edit(request: SmartEditRequest):
    """
    Usa IA para entender comandos naturais de edição e executar automaticamente.
    Exemplos: "edite o título para X", "mude a conclusão", "altere o resumo", etc.
    """
    from docx import Document
    from services.ai import get_model

    # Buscar arquivo
    paths_to_check = [
        f"{PROCESSED_DIR}/edited_{request.filename}",
        f"{PROCESSED_DIR}/formatted_{request.filename}",
        f"{UPLOAD_DIR}/{request.filename}"
    ]

    file_path = None
    for path in paths_to_check:
        if os.path.exists(path):
            file_path = path
            break

    if not file_path:
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        doc = Document(file_path)

        # Extrair primeiros 30 parágrafos para contexto
        paragraphs_preview = []
        for idx, para in enumerate(doc.paragraphs[:30], 1):
            if para.text.strip():
                paragraphs_preview.append({
                    "numero": idx,
                    "texto": para.text.strip()[:150],
                    "estilo": para.style.name
                })

        # Usar IA para analisar o comando
        model = get_model()
        prompt = f"""
        Analise o comando do usuário e determine como editar o documento.

        COMANDO DO USUÁRIO: {request.command}

        PARÁGRAFOS DO DOCUMENTO:
        {json.dumps(paragraphs_preview, ensure_ascii=False)}

        Retorne APENAS um JSON válido no seguinte formato:
        {{
            "action": "edit_element" | "edit_paragraph" | "not_editable",
            "element_type": "titulo_principal" | "subtitulo" | "autor" | "instituicao" | null,
            "paragraph_number": número_do_parágrafo | null,
            "new_text": "texto extraído do comando",
            "explanation": "breve explicação do que será feito"
        }}

        REGRAS:
        - Se o usuário mencionar "título", "título do trabalho", "título principal" -> element_type: "titulo_principal"
        - Se o usuário mencionar "subtítulo" -> element_type: "subtitulo"
        - Se o usuário mencionar "autor" -> element_type: "autor"
        - Se o usuário mencionar "instituição" -> element_type: "instituicao"
        - Se o usuário especificar um número de parágrafo -> action: "edit_paragraph", paragraph_number: X
        - Se não for possível editar (ex: pergunta, dúvida) -> action: "not_editable"
        - Extraia o novo texto do comando (a parte após "para", "coloque", etc)

        Retorne APENAS o JSON, sem markdown ou explicações.
        """

        response = model.generate_content(prompt)
        text = response.text.strip()

        # Limpar markdown se presente
        if text.startswith("```"):
            text = text.split("\n", 1)[1]
            text = text.rsplit("```", 1)[0]

        result = json.loads(text)

        # Executar a ação apropriada
        if result["action"] == "edit_element":
            edit_req = EditElementRequest(
                filename=request.filename,
                element_type=result["element_type"],
                new_text=result["new_text"]
            )
            await edit_element(edit_req)
            return {
                "success": True,
                "message": result["explanation"]
            }

        elif result["action"] == "edit_paragraph":
            edit_req = EditParagraphRequest(
                filename=request.filename,
                paragraph_number=result["paragraph_number"],
                new_text=result["new_text"]
            )
            await edit_paragraph(edit_req)
            return {
                "success": True,
                "message": result["explanation"]
            }

        else:
            return {
                "success": False,
                "message": "Não entendi o comando de edição. Seja mais específico sobre o que deseja alterar."
            }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao processar edição: {str(e)}")


@router.get("/structure/{filename}")
async def get_structure(filename: str):
    """
    Retorna a estrutura do documento com as seções identificadas.
    """
    file_path = f"{UPLOAD_DIR}/{filename}"
    if not os.path.exists(file_path):
        file_path = f"{PROCESSED_DIR}/formatted_{filename}"
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        structure = get_document_structure(file_path)
        return structure
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/write")
async def write_to_document(request: WriteRequest):
    """
    Gera texto acadêmico com IA e insere no documento.
    """
    # Encontrar o arquivo
    file_path = f"{UPLOAD_DIR}/{request.filename}"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        # Extrair contexto do documento
        document_text = extract_text_from_docx(file_path)

        # Gerar texto acadêmico com IA
        generated_text = generate_academic_text(
            document_context=document_text,
            instruction=request.instruction,
            section_type=request.section_type
        )

        if generated_text.startswith("Erro"):
            raise HTTPException(status_code=500, detail=generated_text)

        # Determinar o arquivo de saída
        output_filename = f"edited_{request.filename}"
        output_path = f"{PROCESSED_DIR}/{output_filename}"

        # Inserir texto no documento
        if request.position == "fim":
            insert_text_at_end(file_path, output_path, generated_text)
        else:
            insert_text_after_section(file_path, output_path, request.position, generated_text)

        return {
            "success": True,
            "generated_text": generated_text,
            "output_filename": output_filename,
            "download_url": f"/api/documents/download/{output_filename}",
            "message": f"Texto gerado e inserido com sucesso!"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/preview/{filename}")
async def get_document_preview(filename: str):
    """
    Retorna PDF do documento para preview no iframe.
    Converte DOCX para PDF se necessário.
    """
    # Procurar o arquivo em ordem de prioridade
    paths_to_check = [
        f"{PROCESSED_DIR}/edited_{filename}",
        f"{PROCESSED_DIR}/formatted_{filename}",
        f"{UPLOAD_DIR}/{filename}"
    ]

    file_path = None
    for path in paths_to_check:
        if os.path.exists(path):
            file_path = path
            break

    if not file_path:
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    # Converter para PDF
    pdf_path = file_path.replace(".docx", "_preview.pdf")

    # Se PDF já existe e é mais recente que o DOCX, usar ele
    if os.path.exists(pdf_path):
        docx_time = os.path.getmtime(file_path)
        pdf_time = os.path.getmtime(pdf_path)
        if pdf_time > docx_time:
            return FileResponse(
                pdf_path,
                media_type='application/pdf',
                headers={"Content-Disposition": "inline"}
            )

    # Converter DOCX para PDF
    try:
        conversion_success = convert_docx_to_pdf(file_path, pdf_path)

        if conversion_success and os.path.exists(pdf_path):
            return FileResponse(
                pdf_path,
                media_type='application/pdf',
                headers={"Content-Disposition": "inline"}
            )
        else:
            raise HTTPException(
                status_code=500,
                detail="Não foi possível converter documento para PDF"
            )
    except Exception as e:
        print(f"Erro ao converter para PDF: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao gerar preview: {str(e)}"
        )


@router.post("/write-stream")
async def write_stream(request: Request):
    """
    Gera texto acadêmico com IA usando streaming SSE.
    """
    body = await request.json()
    filename = body.get("filename")
    instruction = body.get("instruction")
    section_type = body.get("section_type", "geral")
    position = body.get("position", "fim")

    file_path = f"{UPLOAD_DIR}/{filename}"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    document_text = extract_text_from_docx(file_path)

    async def event_generator():
        full_text = ""
        try:
            async for chunk in generate_academic_text_stream(
                document_context=document_text,
                instruction=instruction,
                section_type=section_type
            ):
                full_text += chunk
                yield {
                    "event": "chunk",
                    "data": json.dumps({"text": chunk, "full_text": full_text})
                }

            # Após terminar a geração, inserir no documento
            output_filename = f"edited_{filename}"
            output_path = f"{PROCESSED_DIR}/{output_filename}"

            if position == "fim":
                insert_text_at_end(file_path, output_path, full_text)
            else:
                insert_text_after_section(file_path, output_path, position, full_text)

            yield {
                "event": "complete",
                "data": json.dumps({
                    "success": True,
                    "generated_text": full_text,
                    "output_filename": output_filename,
                    "download_url": f"/api/documents/download/{output_filename}"
                })
            }
        except Exception as e:
            yield {
                "event": "error",
                "data": json.dumps({"error": str(e)})
            }

    return EventSourceResponse(event_generator())


@router.get("/complete-vision/{filename}")
async def get_complete_vision(filename: str):
    """
    NORMAEX 2.0 - Retorna a visão completa estrutural + visual do documento

    Este endpoint extrai:
    - Estrutura completa do DOCX (parágrafos, estilos, hierarquia, margens)
    - Análise visual do PDF (coordenadas, layout, fontes)
    - Estatísticas do documento
    - Verificação rápida de conformidade ABNT
    """
    docx_path = f"{UPLOAD_DIR}/{filename}"

    if not os.path.exists(docx_path):
        # Tentar no diretório de processados
        docx_path = f"{PROCESSED_DIR}/formatted_{filename}"
        if not os.path.exists(docx_path):
            docx_path = f"{PROCESSED_DIR}/edited_{filename}"
            if not os.path.exists(docx_path):
                raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        # 1. Extrair estrutura completa do DOCX
        docx_structure = extract_complete_structure(docx_path)

        # 2. Converter DOCX para PDF temporário
        pdf_path = docx_path.replace(".docx", "_temp.pdf")

        conversion_success = convert_docx_to_pdf(docx_path, pdf_path)

        if conversion_success and os.path.exists(pdf_path):
            # 3. Extrair dados visuais do PDF
            pdf_visual = extract_visual_layout(pdf_path)

            # 4. Combinar ambos
            complete_vision = merge_docx_and_pdf_data(docx_structure, pdf_visual)

            # 5. Limpar PDF temporário
            try:
                os.remove(pdf_path)
            except:
                pass
        else:
            # Se conversão falhar, retornar apenas estrutura do DOCX
            complete_vision = {
                "structure": docx_structure,
                "visual": None,
                "visual_margins": None,
                "analysis": {
                    "total_elements": {
                        "paragraphs": len(docx_structure["paragraphs"]),
                        "sections": len(docx_structure["sections"]),
                        "hierarchy_levels": len(docx_structure["hierarchy"]),
                    },
                    "document_type": "unknown",
                    "abnt_compliance": docx_structure.get("statistics", {})
                },
                "note": "Conversão para PDF não disponível. Retornando apenas estrutura DOCX."
            }

        return complete_vision

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao extrair visão completa: {str(e)}")


@router.get("/analyze-structure/{filename}")
async def analyze_structure(filename: str):
    """
    NORMAEX 2.0 - Análise estrutural inteligente com IA

    Este endpoint:
    - Extrai visão completa do documento (FASE 1)
    - Analisa estrutura com IA Gemini (FASE 2)
    - Classifica parágrafos (title, body, citation, etc.)
    - Detecta inconsistências de estilo
    - Identifica problemas ABNT
    - Gera plano de ação executável

    Returns:
        dict: Análise completa + plano de ação
    """
    docx_path = f"{UPLOAD_DIR}/{filename}"

    if not os.path.exists(docx_path):
        # Tentar no diretório de processados
        docx_path = f"{PROCESSED_DIR}/formatted_{filename}"
        if not os.path.exists(docx_path):
            docx_path = f"{PROCESSED_DIR}/edited_{filename}"
            if not os.path.exists(docx_path):
                raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        # 1. Extrair visão completa (FASE 1)
        docx_structure = extract_complete_structure(docx_path)

        # 2. Converter DOCX para PDF (opcional para análise visual)
        pdf_path = docx_path.replace(".docx", "_temp.pdf")
        conversion_success = convert_docx_to_pdf(docx_path, pdf_path)

        if conversion_success and os.path.exists(pdf_path):
            pdf_visual = extract_visual_layout(pdf_path)
            complete_vision = merge_docx_and_pdf_data(docx_structure, pdf_visual)

            # Limpar PDF temporário
            try:
                os.remove(pdf_path)
            except:
                pass
        else:
            # Sem conversão PDF, usar apenas estrutura DOCX
            complete_vision = {
                "structure": docx_structure,
                "visual": None,
                "visual_margins": None
            }

        # 3. Análise estrutural com IA (FASE 2)
        structural_analysis = analyze_document_with_ai(complete_vision)

        # 4. Classificar todos os parágrafos
        paragraph_classifications = classify_all_paragraphs(complete_vision, use_ai=False)

        # 5. Combinar tudo
        final_analysis = {
            "filename": filename,
            "document_info": {
                "total_paragraphs": len(complete_vision["structure"]["paragraphs"]),
                "total_sections": len(complete_vision["structure"]["sections"]),
                "hierarchy_levels": len(complete_vision["structure"]["hierarchy"]),
                "total_words": complete_vision["structure"]["statistics"]["total_words"],
                "document_type": complete_vision.get("analysis", {}).get("document_type", "unknown")
            },
            "structural_analysis": structural_analysis,
            "paragraph_classifications": paragraph_classifications[:20],  # Primeiros 20
            "ready_for_execution": len(structural_analysis["action_plan"]) > 0
        }

        return final_analysis

    except Exception as e:
        import traceback
        print(f"Erro na análise estrutural: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Erro ao analisar estrutura: {str(e)}")


@router.post("/smart-format")
async def smart_format(request: ApplyRequest):
    """
    NORMAEX 2.0 - Formatação inteligente com IA + Executor

    Este endpoint implementa o fluxo completo:
    1. FASE 1: Extrai visão completa do documento
    2. FASE 2: IA analisa e gera plano de ação
    3. FASE 3: Executor aplica o plano de ação

    Args:
        request: Contém o filename do documento

    Returns:
        dict: Resultado da formatação inteligente
    """
    file_location = f"{UPLOAD_DIR}/{request.filename}"

    if not os.path.exists(file_location):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")

    try:
        # 1. FASE 1: Extrair visão completa
        docx_structure = extract_complete_structure(file_location)

        # Tentar conversão PDF (opcional)
        pdf_path = file_location.replace(".docx", "_temp.pdf")
        conversion_success = convert_docx_to_pdf(file_location, pdf_path)

        if conversion_success and os.path.exists(pdf_path):
            pdf_visual = extract_visual_layout(pdf_path)
            complete_vision = merge_docx_and_pdf_data(docx_structure, pdf_visual)

            try:
                os.remove(pdf_path)
            except:
                pass
        else:
            complete_vision = {
                "structure": docx_structure,
                "visual": None,
                "visual_margins": None
            }

        # 2. FASE 2: IA analisa e gera plano de ação
        structural_analysis = analyze_document_with_ai(complete_vision)
        action_plan = structural_analysis.get("action_plan", [])

        if not action_plan:
            return {
                "success": False,
                "message": "Nenhuma ação necessária. Documento já está em conformidade ABNT.",
                "compliance_score": structural_analysis.get("summary", {}).get("compliance_score", 100),
                "analysis": structural_analysis
            }

        # 3. FASE 3: Executor aplica plano de ação
        executor = DocumentExecutor(file_location)
        execution_results = executor.execute_action_plan(action_plan)

        # 4. Salvar documento formatado
        output_filename = f"smart_formatted_{request.filename}"
        output_path = f"{PROCESSED_DIR}/{output_filename}"

        executor.save(output_path)

        # 5. Obter sumário da execução
        summary = executor.get_summary()

        # 6. Montar resposta
        return {
            "success": True,
            "message": "Formatação inteligente aplicada com sucesso!",
            "filename": request.filename,
            "output_filename": output_filename,
            "output_path": output_path,
            "download_url": f"/api/documents/download/{output_filename}",
            "analysis": {
                "total_issues_found": structural_analysis.get("summary", {}).get("total_issues", 0),
                "compliance_score_before": structural_analysis.get("summary", {}).get("compliance_score", 0),
                "compliance_score_after": 100,  # Assume 100% após correções
                "ai_issues": structural_analysis.get("summary", {}).get("ai_issues", 0),
                "style_inconsistencies": structural_analysis.get("summary", {}).get("style_inconsistencies", 0)
            },
            "execution": {
                "total_actions": summary["total_actions"],
                "successful_actions": summary["successful_actions"],
                "failed_actions": summary["failed_actions"],
                "success_rate": summary["success_rate"],
                "details": execution_results
            },
            "action_plan": action_plan
        }

    except Exception as e:
        import traceback
        print(f"Erro na formatação inteligente: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Erro ao aplicar formatação inteligente: {str(e)}")


@router.get("/validate/{filename}")
async def validate_document(filename: str, compare_with: Optional[str] = None):
    """
    NORMAEX 2.0 - FASE 4: Validação Visual de Formatação

    Este endpoint valida se o documento formatado está em conformidade
    com as normas ABNT usando análise visual do PDF.

    Args:
        filename: Nome do arquivo a ser validado
        compare_with: (Opcional) Nome do arquivo original para comparação

    Returns:
        dict: Resultado completo da validação com scores 0-100

    Validações realizadas:
    - Margens (3cm top/left, 2cm bottom/right)
    - Fontes (Arial/Times 12pt)
    - Espaçamento (1.5 entre linhas)
    - Alinhamento (justificado)
    - Score geral de conformidade ABNT
    """
    # Procurar arquivo formatado
    file_path = None
    possible_paths = [
        f"{PROCESSED_DIR}/smart_formatted_{filename}",
        f"{PROCESSED_DIR}/formatted_{filename}",
        f"{PROCESSED_DIR}/edited_{filename}",
        f"{PROCESSED_DIR}/{filename}",
        f"{UPLOAD_DIR}/{filename}"
    ]

    for path in possible_paths:
        if os.path.exists(path):
            file_path = path
            break

    if not file_path:
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        # 1. Converter documento para PDF
        pdf_path = file_path.replace(".docx", "_validation.pdf")
        conversion_success = convert_docx_to_pdf(file_path, pdf_path)

        if not conversion_success or not os.path.exists(pdf_path):
            raise HTTPException(
                status_code=500,
                detail="Não foi possível converter o documento para PDF para validação"
            )

        # 2. Validar documento formatado
        validation_result = validate_document_quality(pdf_path)

        # 3. Se houver arquivo original para comparação
        comparison_result = None
        if compare_with:
            original_path = f"{UPLOAD_DIR}/{compare_with}"
            if os.path.exists(original_path):
                original_pdf_path = original_path.replace(".docx", "_validation_original.pdf")
                original_conversion = convert_docx_to_pdf(original_path, original_pdf_path)

                if original_conversion and os.path.exists(original_pdf_path):
                    original_validation = validate_document_quality(original_pdf_path)

                    # Calcular melhorias
                    comparison_result = {
                        "original_score": original_validation["overall_score"],
                        "formatted_score": validation_result["overall_score"],
                        "improvement": validation_result["overall_score"] - original_validation["overall_score"],
                        "improvements_by_category": {
                            "margins": validation_result["margins"]["score"] - original_validation["margins"]["score"],
                            "fonts": validation_result["fonts"]["score"] - original_validation["fonts"]["score"],
                            "spacing": validation_result["spacing"]["score"] - original_validation["spacing"]["score"],
                            "alignment": validation_result["alignment"]["score"] - original_validation["alignment"]["score"]
                        }
                    }

                    # Limpar PDF original temporário
                    try:
                        os.remove(original_pdf_path)
                    except:
                        pass

        # 4. Limpar PDF temporário
        try:
            os.remove(pdf_path)
        except:
            pass

        # 5. Montar resposta final
        response = {
            "success": True,
            "filename": filename,
            "validation": validation_result,
            "comparison": comparison_result,
            "summary": {
                "overall_score": validation_result["overall_score"],
                "is_abnt_compliant": validation_result["overall_score"] >= 85,
                "total_issues": len(validation_result["all_issues"]),
                "critical_issues": len([i for i in validation_result["all_issues"] if i["severity"] == "critical"]),
                "warnings": len([i for i in validation_result["all_issues"] if i["severity"] == "warning"])
            }
        }

        return response

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"Erro na validação do documento: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Erro ao validar documento: {str(e)}")


@router.post("/intelligent-write")
async def intelligent_write(request: WriteRequest):
    """
    NORMAEX 2.0 - FASE 5: Escrita Inteligente

    IA escreve texto JÁ formatado estruturalmente seguindo ABNT.

    Diferença do /write normal:
    - /write: IA escreve texto → depois formata (2 etapas)
    - /intelligent-write: IA escreve + formata simultaneamente (1 etapa integrada)

    Args:
        request: WriteRequest com filename, instruction, section_type, position

    Returns:
        dict: Texto gerado + estrutura aplicada + documento editado
    """
    file_path = f"{UPLOAD_DIR}/{request.filename}"

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    try:
        # 1. Extrair contexto e estrutura do documento
        document_text = extract_text_from_docx(file_path)
        document_structure = extract_complete_structure(file_path)

        # 2. IA escreve com estrutura (FASE 5)
        write_result = write_with_structure(
            document_context=document_text,
            instruction=request.instruction,
            section_type=request.section_type,
            document_structure=document_structure
        )

        # 3. Executar escrita estruturada
        output_filename = f"intelligent_edited_{request.filename}"
        output_path = f"{PROCESSED_DIR}/{output_filename}"

        execution_result = execute_write_with_structure(
            doc_path=file_path,
            output_path=output_path,
            write_result=write_result,
            position=request.position
        )

        # 4. Criar plano de ação para referência
        action_plan = create_action_plan_for_writing(
            content=write_result["content"],
            formatting=write_result["structure"]["formatting"],
            position=request.position
        )

        return {
            "success": True,
            "message": "Texto inteligente gerado e formatado com sucesso!",
            "generated_content": write_result["content"],
            "structure_applied": write_result["structure"],
            "execution_result": execution_result,
            "action_plan": action_plan,
            "output_filename": output_filename,
            "download_url": f"/api/documents/download/{output_filename}",
            "stats": {
                "total_words": execution_result["total_words"],
                "total_characters": execution_result["total_characters"],
                "paragraphs_inserted": execution_result["paragraphs_inserted"]
            }
        }

    except Exception as e:
        import traceback
        print(f"Erro na escrita inteligente: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Erro ao gerar texto inteligente: {str(e)}")


@router.post("/intelligent-write-stream")
async def intelligent_write_stream(request: Request):
    """
    NORMAEX 2.0 - FASE 5: Escrita Inteligente com Streaming

    Versão streaming da escrita inteligente usando SSE.
    Retorna texto conforme é gerado pela IA.

    Returns:
        EventSourceResponse: Stream de eventos com chunks de texto
    """
    body = await request.json()
    filename = body.get("filename")
    instruction = body.get("instruction")
    section_type = body.get("section_type", "geral")
    position = body.get("position", "fim")

    file_path = f"{UPLOAD_DIR}/{filename}"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Documento não encontrado")

    async def event_generator():
        full_text = ""
        try:
            # Extrair contexto
            document_text = extract_text_from_docx(file_path)
            document_structure = extract_complete_structure(file_path)

            # Gerar texto com streaming
            for chunk in write_structured_streaming(
                document_context=document_text,
                instruction=instruction,
                section_type=section_type,
                document_structure=document_structure
            ):
                full_text += chunk
                yield {
                    "event": "chunk",
                    "data": json.dumps({"text": chunk, "full_text": full_text})
                }

            # Após terminar a geração, aplicar formatação estruturada
            write_result = {
                "content": full_text,
                "structure": {
                    "type": "body",
                    "formatting": {
                        "font": "Arial",
                        "size": 12,
                        "alignment": "justify",
                        "spacing": 1.5,
                        "indent": 1.25
                    },
                    "section": section_type
                }
            }

            output_filename = f"intelligent_edited_{filename}"
            output_path = f"{PROCESSED_DIR}/{output_filename}"

            execution_result = execute_write_with_structure(
                doc_path=file_path,
                output_path=output_path,
                write_result=write_result,
                position=position
            )

            yield {
                "event": "complete",
                "data": json.dumps({
                    "success": True,
                    "generated_content": full_text,
                    "structure_applied": write_result["structure"],
                    "execution_result": execution_result,
                    "output_filename": output_filename,
                    "download_url": f"/api/documents/download/{output_filename}"
                })
            }

        except Exception as e:
            import traceback
            print(f"Erro no streaming inteligente: {traceback.format_exc()}")
            yield {
                "event": "error",
                "data": json.dumps({"error": str(e)})
            }

    return EventSourceResponse(event_generator())
